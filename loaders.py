"""
loaders.py
==========
One function per supported file type, plus a dispatcher (`load_file`) that
routes to the right loader based on extension and guarantees the caller
never sees an exception -- a bad file returns an empty list and gets logged,
it never takes down the whole ingestion run.

Every loader returns list[Document] with a *base* metadata set
(source_file, filename, folder, extension, company, created_at, page).
splitter.py later adds 'chunk_number' and 'document_type' when it splits
each Document into chunks.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Callable

import pandas as pd
from langchain_core.documents import Document

import config
import utils

logger = utils.logger


# ---------------------------------------------------------------------------
# Shared metadata builder
# ---------------------------------------------------------------------------
def _base_metadata(path: Path, page: int | None = None) -> dict:
    """Build the minimum required metadata dict shared by every loader."""
    return {
        "source_file": str(path),
        "filename": path.name,
        "folder": str(path.parent),
        "extension": path.suffix.lower(),
        "page": page,
        "company": utils.infer_company(path),
        "created_at": utils.utc_now_iso(),
    }


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------
def load_csv(path: Path) -> list[Document]:
    """
    Load a CSV file, one Document per row.

    Handles: missing columns, encoding issues, malformed CSV, empty files.
    Falls back through a list of encodings before giving up.
    """
    encodings_to_try = ("utf-8", "utf-8-sig", "latin-1", "cp1252")
    df: pd.DataFrame | None = None
    last_error: Exception | None = None

    for encoding in encodings_to_try:
        try:
            df = pd.read_csv(path, encoding=encoding, on_bad_lines="skip")
            break
        except (UnicodeDecodeError, pd.errors.ParserError, OSError) as exc:
            last_error = exc
            continue

    if df is None:
        logger.error("CSV load failed for %s after trying all encodings: %s", path, last_error)
        return []

    if df.empty:
        logger.warning("CSV file %s is empty; skipping.", path)
        return []

    documents: list[Document] = []
    for row_idx, row in df.iterrows():
        # Join all non-null column values into the page content. This keeps
        # the loader generic instead of assuming fixed column names like
        # 'Title'/'Review', which only fit one particular dataset.
        try:
            text_parts = [str(v) for v in row.values if pd.notna(v)]
            content = " ".join(text_parts).strip()
        except Exception as exc:  # defensive: malformed row values
            logger.warning("Skipping malformed row %s in %s: %s", row_idx, path, exc)
            continue

        if not content:
            continue

        metadata = _base_metadata(path)
        metadata["row_index"] = int(row_idx)
        # Include original columns as metadata where they are scalar-safe.
        for col, val in row.to_dict().items():
            metadata[f"col_{col}"] = utils.safe_str(val) if pd.notna(val) else None

        documents.append(Document(page_content=content, metadata=metadata))

    return documents


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def load_pdf(path: Path) -> list[Document]:
    """
    Load a PDF, one Document per page, using pypdf directly (avoids pulling
    in the heavier 'unstructured' PDF stack for the common case).

    Guards against malicious / huge PDFs via MAX_PDF_PAGES.
    """
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError:
        logger.error("pypdf is not installed; cannot load PDF %s", path)
        return []

    try:
        reader = PdfReader(str(path))
    except (PdfReadError, OSError, ValueError) as exc:
        logger.error("Unreadable/corrupted PDF %s: %s", path, exc)
        return []

    if reader.is_encrypted:
        try:
            reader.decrypt("")  # try an empty password; otherwise give up safely
        except Exception:
            logger.warning("Encrypted PDF %s could not be opened; skipping.", path)
            return []

    num_pages = len(reader.pages)
    if num_pages == 0:
        logger.warning("PDF %s has no pages; skipping.", path)
        return []

    if num_pages > config.MAX_PDF_PAGES:
        logger.warning(
            "PDF %s has %d pages, exceeding MAX_PDF_PAGES=%d; truncating.",
            path, num_pages, config.MAX_PDF_PAGES,
        )
        num_pages = config.MAX_PDF_PAGES

    documents: list[Document] = []
    for page_number in range(num_pages):
        try:
            text = reader.pages[page_number].extract_text() or ""
        except Exception as exc:  # pypdf can raise on malformed page objects
            logger.warning("Could not extract page %d of %s: %s", page_number, path, exc)
            continue

        text = text.strip()
        if not text:
            continue

        metadata = _base_metadata(path, page=page_number + 1)
        documents.append(Document(page_content=text, metadata=metadata))

    return documents


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------
def load_txt(path: Path) -> list[Document]:
    """Load a plain text file, trying a few encodings before giving up."""
    encodings_to_try = ("utf-8", "utf-8-sig", "latin-1", "cp1252")

    for encoding in encodings_to_try:
        try:
            text = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, OSError) as exc:
            last_error = exc
            continue
    else:
        logger.error("TXT load failed for %s: %s", path, last_error)
        return []

    text = text.strip()
    if not text:
        logger.warning("TXT file %s is empty; skipping.", path)
        return []

    return [Document(page_content=text, metadata=_base_metadata(path))]


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def load_docx(path: Path) -> list[Document]:
    """
    Load a Word document's paragraph text using python-docx.

    python-docx never executes macros or embedded content, so this is safe
    against "malicious document" style attacks by construction.
    """
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed; cannot load DOCX %s", path)
        return []

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # zipfile.BadZipFile, docx.opc.exceptions, etc.
        logger.error("Corrupted or unreadable DOCX %s: %s", path, exc)
        return []

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs).strip()

    if not text:
        logger.warning("DOCX file %s has no extractable text; skipping.", path)
        return []

    return [Document(page_content=text, metadata=_base_metadata(path))]


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------
def load_markdown(path: Path) -> list[Document]:
    """
    Load a Markdown file as plain text (headings/structure are preserved as
    literal text, which RecursiveCharacterTextSplitter handles well enough
    for retrieval purposes without pulling in a full Markdown AST parser).
    """
    docs = load_txt(path)  # markdown is just text with a different extension
    for doc in docs:
        doc.metadata["extension"] = path.suffix.lower()
    return docs


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------
def load_json(path: Path) -> list[Document]:
    """
    Load a JSON file.

    * If the top-level structure is a list of objects, one Document per
      element (common for "records" style exports).
    * Otherwise, the whole file becomes a single Document.

    Guards against pathological JSON via MAX_JSON_RECORDS.
    """
    try:
        raw_text = path.read_text(encoding="utf-8")
    except (UnicodeDecodeError, OSError) as exc:
        logger.error("Could not read JSON file %s: %s", path, exc)
        return []

    try:
        data = json.loads(raw_text)
    except json.JSONDecodeError as exc:
        logger.error("Corrupted JSON in %s: %s", path, exc)
        return []

    documents: list[Document] = []

    if isinstance(data, list):
        if len(data) > config.MAX_JSON_RECORDS:
            logger.warning(
                "JSON file %s has %d records, exceeding MAX_JSON_RECORDS=%d; truncating.",
                path, len(data), config.MAX_JSON_RECORDS,
            )
            data = data[: config.MAX_JSON_RECORDS]

        for idx, record in enumerate(data):
            content = json.dumps(record, ensure_ascii=False) if not isinstance(record, str) else record
            content = content.strip()
            if not content:
                continue
            metadata = _base_metadata(path)
            metadata["record_index"] = idx
            documents.append(Document(page_content=content, metadata=metadata))
    else:
        content = json.dumps(data, ensure_ascii=False, indent=2).strip()
        if content:
            documents.append(Document(page_content=content, metadata=_base_metadata(path)))

    if not documents:
        logger.warning("JSON file %s produced no documents; skipping.", path)

    return documents


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
def load_excel(path: Path) -> list[Document]:
    """
    Load every sheet of an Excel workbook, one Document per row
    (mirrors the CSV loader). Invalid/unreadable sheets are skipped
    individually rather than failing the whole workbook.
    """
    try:
        sheets = pd.read_excel(path, sheet_name=None, engine=None)
    except (ValueError, OSError, Exception) as exc:  # openpyxl raises various types
        logger.error("Could not open Excel file %s: %s", path, exc)
        return []

    documents: list[Document] = []

    for sheet_name, df in sheets.items():
        if df.empty:
            logger.warning("Sheet '%s' in %s is empty; skipping.", sheet_name, path)
            continue

        if len(df) > config.MAX_EXCEL_ROWS:
            logger.warning(
                "Sheet '%s' in %s has %d rows, exceeding MAX_EXCEL_ROWS=%d; truncating.",
                sheet_name, path, len(df), config.MAX_EXCEL_ROWS,
            )
            df = df.head(config.MAX_EXCEL_ROWS)

        for row_idx, row in df.iterrows():
            try:
                text_parts = [str(v) for v in row.values if pd.notna(v)]
                content = " ".join(text_parts).strip()
            except Exception as exc:
                logger.warning(
                    "Skipping malformed row %s in sheet '%s' of %s: %s",
                    row_idx, sheet_name, path, exc,
                )
                continue

            if not content:
                continue

            metadata = _base_metadata(path)
            metadata["sheet_name"] = sheet_name
            metadata["row_index"] = int(row_idx)
            documents.append(Document(page_content=content, metadata=metadata))

    return documents


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
LOADER_MAP: dict[str, Callable[[Path], list[Document]]] = {
    ".csv": load_csv,
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".docx": load_docx,
    ".md": load_markdown,
    ".markdown": load_markdown,
    ".json": load_json,
    ".xlsx": load_excel,
    ".xls": load_excel,
}


def load_file(path: Path) -> list[Document]:
    """
    Single entry point used by ingest.py.

    Performs all safety checks (path sanitization, extension support, size
    limit) before dispatching to the type-specific loader, and guarantees
    that no exception ever propagates out of this function -- a bad file
    results in an empty list plus a logged error, never a crash.
    """
    try:
        safe_path = utils.sanitize_path(path)
    except utils.UnsafePathError as exc:
        logger.error("Rejected unsafe path %s: %s", path, exc)
        return []

    if not utils.is_supported_extension(safe_path):
        logger.debug("Skipping unsupported file type: %s", safe_path)
        return []

    if not utils.check_file_size(safe_path):
        logger.warning(
            "Skipping %s: file is empty or exceeds MAX_FILE_SIZE_MB=%d.",
            safe_path, config.MAX_FILE_SIZE_MB,
        )
        return []

    loader_fn = LOADER_MAP.get(safe_path.suffix.lower())
    if loader_fn is None:
        logger.debug("No loader registered for extension %s", safe_path.suffix)
        return []

    try:
        documents = loader_fn(safe_path)
    except PermissionError as exc:
        logger.error("Permission denied reading %s: %s", safe_path, exc)
        return []
    except Exception as exc:  # final safety net -- one bad file must never kill the run
        logger.error("Unexpected error loading %s: %s", safe_path, exc)
        return []

    logger.info("Loaded %d document(s) from %s", len(documents), safe_path)
    return documents
